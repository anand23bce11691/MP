from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"D:\MP")
OUT = ROOT / "output" / "IncidentIQ_System_and_Database_Design.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "F2F4F7"
WIDTHS = {2: [4680, 4680], 3: [2200, 3600, 3560], 4: [1500, 2200, 2900, 2760]}

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    widths = WIDTHS.get(cols, [9360 // cols] * cols)
    for r_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for c_index, value in enumerate(row_values):
            cells[c_index].text = value.strip()
            if r_index == 0:
                set_cell_shading(cells[c_index], GRAY)
                for run in cells[c_index].paragraphs[0].runs:
                    run.bold = True
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.line_spacing = 1
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(8)

def parse_markdown(doc, text):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            fence = line[3:]
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            if fence == "mermaid":
                doc.add_paragraph("Architecture / relationship diagram definition", style="Heading 3")
            add_code(doc, block)
        elif re.match(r"^#{1,3} ", line):
            level = len(line) - len(line.lstrip("#"))
            doc.add_paragraph(line[level + 1:].strip(), style=f"Heading {level}")
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[-| :]+\|$", lines[i]):
                    rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue
        elif re.match(r"^- ", line):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line).strip(), style="List Number")
        elif not line.strip():
            pass
        else:
            doc.add_paragraph(line.strip())
        i += 1

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "IncidentIQ | System and Database Design"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)
footer = section.footer.paragraphs[0]
footer.text = "IncidentIQ design package"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(72)
title.paragraph_format.space_after = Pt(8)
r = title.add_run("IncidentIQ")
r.font.name = "Calibri"
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(20)
r = subtitle.add_run("System Design, Database Design, and SQL Server Schema")
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(89, 89, 89)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prepared project implementation package").italic = True
doc.add_page_break()

sections = [
    ("Part A - System Design", ROOT / "docs" / "SYSTEM_DESIGN.md"),
    ("Part B - Database Design", ROOT / "docs" / "DATABASE_DESIGN.md"),
]
for index, (label, source) in enumerate(sections):
    doc.add_paragraph(label, style="Heading 1")
    parse_markdown(doc, source.read_text(encoding="utf-8"))
    doc.add_page_break()

doc.add_paragraph("Part C - SQL Server Initial Schema", style="Heading 1")
doc.add_paragraph("The following executable script creates the IncidentIQ SQL Server database tables, constraints, and indexes.")
sql_lines = (ROOT / "database" / "001_initial_schema.sql").read_text(encoding="utf-8").splitlines()
for start in range(0, len(sql_lines), 42):
    add_code(doc, sql_lines[start:start + 42])
    if start + 42 < len(sql_lines):
        doc.add_page_break()

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
